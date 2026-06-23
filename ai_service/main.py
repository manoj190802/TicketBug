import io
import re
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import pypdf
import docx

app = FastAPI(title="TicketBug AI Classifier API")

# Enable CORS so backend/frontend can call it if needed
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Lists of keywords for classification and skill extraction
FRONTEND_KEYWORDS = {
    'angular', 'react', 'vue', 'html', 'css', 'typescript', 'javascript', 'js', 'ts', 
    'scss', 'tailwind', 'bootstrap', 'redux', 'rxjs', 'flexbox', 'responsive', 'ui', 'ux', 'frontend'
}

BACKEND_KEYWORDS = {
    'dotnet', 'c#', 'python', 'fastapi', 'django', 'flask', 'java', 'spring', 'node', 'express', 
    'mongodb', 'sql', 'postgres', 'mysql', 'database', 'db', 'rest api', 'graphql', 'docker', 
    'aws', 'azure', 'gcp', 'microservices', 'backend', 'c++', 'rust', 'go', 'golang'
}

def extract_text_from_pdf(content: bytes) -> str:
    pdf_file = io.BytesIO(content)
    reader = pypdf.PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def extract_text_from_docx(content: bytes) -> str:
    docx_file = io.BytesIO(content)
    doc = docx.Document(docx_file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"
    return text

@app.post("/analyze")
async def analyze_requirements(file: UploadFile = File(...)):
    filename = file.filename.lower()
    content = await file.read()
    
    text = ""
    try:
        if filename.endswith(".pdf"):
            text = extract_text_from_pdf(content)
        elif filename.endswith(".docx"):
            text = extract_text_from_docx(content)
        elif filename.endswith(".txt") or filename.endswith(".json") or filename.endswith(".md"):
            try:
                text = content.decode("utf-8")
            except UnicodeDecodeError:
                text = content.decode("latin-1")
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Please upload .txt, .pdf, or .docx.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse document: {str(e)}")

    if not text.strip():
        raise HTTPException(status_code=400, detail="The uploaded document is empty or could not be parsed.")

    # Lowercase text for matching
    text_lower = text.lower()
    
    # Simple word extraction using alphanumeric split
    words = set(re.findall(r'\b[a-z0-9#+.-]+\b', text_lower))
    
    # Find matching keywords
    matched_frontend = FRONTEND_KEYWORDS.intersection(words)
    matched_backend = BACKEND_KEYWORDS.intersection(words)
    
    # Count occurrences of frontend and backend keywords
    fe_count = sum(1 for w in re.findall(r'\b[a-z0-9#+.-]+\b', text_lower) if w in FRONTEND_KEYWORDS)
    be_count = sum(1 for w in re.findall(r'\b[a-z0-9#+.-]+\b', text_lower) if w in BACKEND_KEYWORDS)
    
    # Add multi-word checks
    if 'rest api' in text_lower or 'restful api' in text_lower or 'web api' in text_lower:
        matched_backend.add('rest api')
        be_count += 2
        
    # Skills extracted are the union of matched keywords
    skills = list(matched_frontend.union(matched_backend))
    
    # Classify category
    if len(matched_frontend) >= 2 and len(matched_backend) >= 2:
        category = "Full Stack"
    elif fe_count > be_count:
        category = "Frontend"
    elif be_count > fe_count:
        category = "Backend"
    elif len(matched_frontend) > 0 and len(matched_backend) == 0:
        category = "Frontend"
    elif len(matched_backend) > 0 and len(matched_frontend) == 0:
        category = "Backend"
    else:
        category = "Full Stack" # Default fallback
        
    return {
        "fileName": file.filename,
        "category": category,
        "extractedSkills": skills,
        "summary": text[:200] + "..." if len(text) > 200 else text,
        "wordCount": len(text.split())
    }

@app.get("/health")
def health_check():
    return {"status": "healthy"}
